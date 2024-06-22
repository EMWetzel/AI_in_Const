
#Activate Conda QA_Env for run
#conda activate QA_ENV

#Activate Conda QA_Env for run
#conda activate QA_ENV

from transformers import pipeline, T5Tokenizer, T5ForConditionalGeneration
from docx import Document

# Load models and tokenizers
question_generation_model_name = "valhalla/t5-small-qg-hl"
answer_generation_model_name = "t5-base"

question_tokenizer = T5Tokenizer.from_pretrained(question_generation_model_name)
question_model = T5ForConditionalGeneration.from_pretrained(question_generation_model_name)

answer_tokenizer = T5Tokenizer.from_pretrained(answer_generation_model_name)
answer_model = T5ForConditionalGeneration.from_pretrained(answer_generation_model_name)

summarizer = pipeline("summarization", model="facebook/bart-large-cnn")

def generate_qa_pairs(text):
    inputs = question_tokenizer.encode("highlight: " + text + " </s>", return_tensors="pt")
    outputs = question_model.generate(inputs, max_length=150, num_return_sequences=5, do_sample=True)

    questions = [question_tokenizer.decode(output, skip_special_tokens=True) for output in outputs]
    
    qa_pairs = []
    for question in questions:
        input_text = "question: " + question + " context: " + text + " </s>"
        inputs = answer_tokenizer.encode(input_text, return_tensors="pt")
        outputs = answer_model.generate(inputs, max_length=150, num_return_sequences=1, do_sample=True)
        
        answer = answer_tokenizer.decode(outputs[0], skip_special_tokens=True)
        qa_pairs.append({"question": question, "answer": answer})
    
    return qa_pairs

def save_to_word(qa_pairs, output_file):
    document = Document()
    document.add_heading('Generated Q&A Pairs', 0)
    
    for i, pair in enumerate(qa_pairs):
        document.add_heading(f'Q{i+1}: {pair["question"]}', level=1)
        document.add_paragraph(pair["answer"])
    
    document.save(output_file)

def read_text_from_file(file_path):
    try:
        with open(file_path, 'r', encoding='utf-8') as file:
            text = file.read()
        print("File read successfully!")
        return text
    except Exception as e:
        print(f"An error occurred: {e}")
        return None

# New function to tokenize the document and print tokens
def tokenize_text(text):
    tokens = question_tokenizer.tokenize(text)
    print("Tokenized text:", tokens)

# Read text from a file
input_file = "C:/Users/emw0009/Box/01 - CARV Lab/Research Code/QA Gen/TestText.txt"  # Replace with your .txt file path
text = read_text_from_file(input_file)

if text:
    # Tokenize the text and print tokens
    tokenize_text(text)

    # Generate Q&A pairs
    qa_pairs = generate_qa_pairs(text)

    # Save Q&A pairs to Word document
    output_file = "qa_pairs.docx"
    save_to_word(qa_pairs, output_file)

    print(f"Q&A pairs saved to {output_file}")

